import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import datetime

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="AI RPM Generator", layout="wide")

# --- HEADER & LOGO ---
st.title("🤖 AI Generator Rencana Pembelajaran Mendalam (RPM)")
uploaded_logo = st.file_uploader("Unggah Logo Sekolah (JPG/PNG)", type=["jpg", "png", "jpeg"])

if uploaded_logo:
    st.image(uploaded_logo, width=100)

# --- INPUT FORM ---
with st.form("form_rpm"):
    st.header("📋 Format Input Data")
    
    col1, col2 = st.columns(2)
    with col1:
        satuan_pendidikan = st.text_input("Nama Satuan Pendidikan")
        nama_guru = st.text_input("Nama Guru")
        nip_guru = st.text_input("NIP Guru")
        nama_kepsek = st.text_input("Nama Kepala Sekolah")
        jenjang = st.selectbox("Jenjang Pendidikan", ["SD", "SMP", "SMA"])
        fase = st.selectbox("Fase", ["A", "B", "C", "D", "E", "F"])
        kelas = st.text_input("Kelas")
        semester = st.radio("Semester", ["Ganjil", "Genap"])
        mapel = st.text_input("Mata Pelajaran")
        materi = st.text_area("Materi Pelajaran")
        
    with col2:
        cp_fase = st.text_area("CP Per Fase")
        cp_elemen = st.text_area("CP Per Elemen")
        jml_pertemuan = st.number_input("Jumlah Pertemuan (Angka)", min_value=1, step=1)
        alokasi_waktu = st.text_input("Alokasi Waktu (Misal: 2x45 Menit)")
        
        st.write("**Praktik Pedagogis**")
        pendekatan = st.text_input("Pendekatan Pembelajaran")
        model_pemb = st.text_input("Model Pembelajaran (Misal: PjBL, PBL, Discovery)")
        metode = st.text_input("Metode")
        teknik = st.text_input("Teknik")
        
        dimensi = st.multiselect("Dimensi Lulusan", 
            ["Keimanan dan ketaqwaan", "Kewargaan", "Penalaran kritis", "Kreativitas", "Kolaborasi", "Kemandirian", "Kesehatan", "Komunikasi"])
        
    st.divider()
    tujuan_pemb = st.text_area("Tujuan Pembelajaran (Input awal)")
    disiplin_ilmu = st.text_input("Disiplin Ilmu")
    kemitraan = st.text_input("Kemitraan Pembelajaran")
    lingkungan = st.text_input("Lingkungan Pembelajaran")
    digital = st.text_input("Pemanfaatan Digital")
    sarpras = st.text_area("Sarana Dan Prasarana")
    tanggal_buat = st.date_input("Tanggal Pembuatan", datetime.date.today())

    submitted = st.form_submit_button("Generate RPM via AI")

# --- LOGIKA GENERATOR (SIMULASI AI) ---
def generate_rpm_content(data):
    # Di sini Anda biasanya memanggil API OpenAI/Claude. 
    # Sebagai contoh, saya buatkan struktur string berdasarkan input.
    
    content = f"""
    1. IDENTIFIKASI
    A. Peserta Didik
    - Aspek Kognitif: Mampu memahami konsep {data['materi']} dengan level HOTS.
    - Aspek Sikap: Menunjukkan dimensi {', '.join(data['dimensi'])}.
    - Aspek Keterampilan: Terampil dalam praktik {data['mapel']}.
    B. Materi: {data['materi']}
    C. Dimensi Lulusan: {', '.join(data['dimensi'])}

    2. DESAIN PEMBELAJARAN
    A. CP: {data['cp_fase']}
    B. Tujuan (ABCD): Melalui model {data['model_pemb']}, peserta didik (A) mampu mendemonstrasikan (B) materi {data['materi']} (C) dengan tepat (D).
    C. Disiplin Ilmu: {data['disiplin_ilmu']}
    D. Praktik Pedagogis: Pendekatan {data['pendekatan']}, Model {data['model_pemb']}
    E. Kemitraan: {data['kemitraan']}
    ... (dan seterusnya sesuai output yang diminta)
    """
    return content

# --- FUNGSI DOWNLOAD WORD ---
def create_word(content, data):
    doc = Document()
    doc.add_heading('RENCANA PEMBELAJARAN MENDALAM (RPM)', 0)
    doc.add_paragraph(f"Satuan Pendidikan: {data['satuan_pendidikan']}")
    doc.add_paragraph(content)
    
    # Footer Tanda Tangan
    table = doc.add_table(rows=1, cols=2)
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    
    left_cell.text = f"\n\nMengetahui,\nKepala Sekolah\n\n\n\n{data['nama_kepsek']}\nNIP. {data['nip_kepsek']}"
    right_cell.text = f"{data['tgl']}\nGuru Mata Pelajaran\n\n\n\n{data['nama_guru']}\nNIP. {data['nip_guru']}"
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- DISPLAY HASIL ---
if submitted:
    st.info("AI sedang menyusun RPM terstruktur... (Sintak model & Asesmen sedang dibuat)")
    
    # Gabungkan data input
    form_data = {
        "satuan_pendidikan": satuan_pendidikan, "nama_guru": nama_guru, "nip_guru": nip_guru,
        "nama_kepsek": nama_kepsek, "nip_kepsek": "123456789", # contoh
        "dimensi": dimensi, "materi": materi, "model_pemb": model_pemb,
        "cp_fase": cp_fase, "tgl": tanggal_buat, "mapel": mapel, "disiplin_ilmu": disiplin_ilmu,
        "kemitraan": kemitraan, "pendekatan": pendekatan
    }
    
    # Eksekusi Generator
    hasil_generate = generate_rpm_content(form_data)
    
    st.subheader("📄 Hasil Generate RPM")
    st.text_area("Draft Hasil (Bisa Disalin)", hasil_generate, height=400)
    
    # Tombol Unduh
    col_dl1, col_dl2, col_dl3 = st.columns(3)
    
    word_file = create_word(hasil_generate, form_data)
    with col_dl1:
        st.download_button("📥 Unduh .DOCX (Word)", data=word_file, file_name="RPM_Sekolah.docx")
    
    with col_dl2:
        st.button("📋 Salin Teks (Klik Kanan Copy)")

    # Riwayat Download (Simulasi)
    if 'history' not in st.session_state:
        st.session_state['history'] = []
    st.session_state['history'].append(f"RPM_{mapel}_{tanggal_buat}.docx")

st.sidebar.header("📂 Hasil Download-an (Sesi Ini)")
for h in st.session_state.get('history', []):
    st.sidebar.write(f"✅ {h}")
